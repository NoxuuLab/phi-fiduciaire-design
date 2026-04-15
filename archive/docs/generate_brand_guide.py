"""
PHI Fiduciaire Brand Style Guide Generator
Generates a professional .docx brand style guide document.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Brand Colors ────────────────────────────────────────────────────────────
TEAL        = RGBColor(0x1b, 0x3f, 0x3f)   # Primary teal  #1b3f3f
TEAL_DARK   = RGBColor(0x0f, 0x26, 0x26)   # Footer dark   #0f2626
RED         = RGBColor(0xef, 0x3a, 0x24)   # Accent red    #ef3a24
WHITE       = RGBColor(0xff, 0xff, 0xff)
OFF_WHITE   = RGBColor(0xf8, 0xf6, 0xf2)
DARK_TEXT   = RGBColor(0x1a, 0x1a, 0x1a)
MID_GREY    = RGBColor(0x55, 0x55, 0x55)
LIGHT_GREY  = RGBColor(0xee, 0xee, 0xee)
TEAL_HEX    = "1B3F3F"
TEAL_DARK_HEX = "0F2626"
RED_HEX     = "EF3A24"
OFF_WHITE_HEX = "F8F6F2"
LIGHT_GREY_HEX = "EEEEEE"
WHITE_HEX   = "FFFFFF"

OUTPUT_PATH = r"C:\Users\ledbu\Documents\CLIENTS\PHI FIDUCIAIRE\phi-fiduciaire-design\PHI-Fiduciaire-Brand-Style-Guide.docx"


# ─── XML Helpers ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_para_border_bottom(para, color_hex, sz=12):
    """Add a bottom border to a paragraph (used as a divider line)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_table_borders(table, color_hex="CCCCCC", sz=4):
    """Set all borders on a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color_hex)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_column_width(cell, width_cm):
    """Set column width in a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # 567 twips per cm
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    """Set internal cell padding."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def set_vertical_align(cell, align='center'):
    """Set vertical alignment in table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def add_bookmark(para, bookmark_id, name):
    """Add a bookmark to a paragraph for TOC linking."""
    run = para.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), name)
    tag.append(start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    tag.append(end)


def set_spacing(para, before=0, after=0, line=None, line_rule=None):
    """Set paragraph spacing."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
    if line_rule:
        spacing.set(qn('w:lineRule'), line_rule)
    pPr.append(spacing)


# ─── Document Helpers ─────────────────────────────────────────────────────────

def add_page_break(doc):
    """Add a page break paragraph."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    set_spacing(p, 0, 0)


def heading1(doc, text, bm_id=None, bm_name=None):
    """Add a Heading 1 styled paragraph."""
    para = doc.add_paragraph()
    set_spacing(para, before=280, after=100)
    # Red accent bar above
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top_b = OxmlElement('w:top')
    top_b.set(qn('w:val'), 'single')
    top_b.set(qn('w:sz'), '6')
    top_b.set(qn('w:space'), '4')
    top_b.set(qn('w:color'), RED_HEX)
    pBdr.append(top_b)
    pPr.append(pBdr)

    run = para.add_run(text)
    run.font.name = 'Garamond'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = TEAL
    if bm_id is not None and bm_name:
        add_bookmark(para, bm_id, bm_name)
    return para


def heading2(doc, text):
    """Add a Heading 2 styled paragraph."""
    para = doc.add_paragraph()
    set_spacing(para, before=200, after=80)
    run = para.add_run(text)
    run.font.name = 'Garamond'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = TEAL
    return para


def heading3(doc, text):
    """Add a Heading 3 styled paragraph."""
    para = doc.add_paragraph()
    set_spacing(para, before=160, after=60)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK_TEXT
    return para


def body(doc, text, italic=False, color=None):
    """Add a body text paragraph."""
    para = doc.add_paragraph()
    set_spacing(para, before=40, after=80)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.italic = italic
    run.font.color.rgb = color if color else DARK_TEXT
    return para


def note(doc, text):
    """Add a note/caption paragraph."""
    para = doc.add_paragraph()
    set_spacing(para, before=40, after=60)
    # Light left bar
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), '12')
    left_b.set(qn('w:space'), '6')
    left_b.set(qn('w:color'), RED_HEX)
    pBdr.append(left_b)
    pPr.append(pBdr)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '200')
    pPr.append(ind)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = MID_GREY
    return para


def label(doc, text):
    """Add an eyebrow/label paragraph (UPPERCASE, small, teal)."""
    para = doc.add_paragraph()
    set_spacing(para, before=160, after=40)
    run = para.add_run(text.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = TEAL
    # letter spacing via rPr
    rPr = run._r.get_or_add_rPr()
    spacing_el = OxmlElement('w:spacing')
    spacing_el.set(qn('w:val'), '60')
    rPr.append(spacing_el)
    return para


def divider(doc):
    """Add a thin teal horizontal rule paragraph."""
    para = doc.add_paragraph()
    set_spacing(para, before=100, after=100)
    set_para_border_bottom(para, TEAL_HEX, sz=4)
    return para


def bullet(doc, text, level=0):
    """Add a bullet list item."""
    para = doc.add_paragraph(style='List Bullet')
    set_spacing(para, before=30, after=30)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_TEXT
    return para


def make_table(doc, headers, rows, col_widths_cm=None, header_bg=TEAL_HEX, stripe=True):
    """
    Create a styled table with headers and rows.
    headers: list of strings
    rows: list of list of strings
    col_widths_cm: list of floats (cm), must sum to ~16.5
    """
    ncols = len(headers)
    if col_widths_cm is None:
        col_widths_cm = [16.5 / ncols] * ncols

    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color_hex="CCCCCC", sz=4)

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        set_column_width(cell, col_widths_cm[i])
        set_vertical_align(cell, 'center')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "F2F8F8" if stripe and ri % 2 == 0 else WHITE_HEX
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            set_column_width(cell, col_widths_cm[ci])
            set_vertical_align(cell, 'center')
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_TEXT

    return table


def add_color_swatch_row(doc, swatches):
    """
    Add a row of color swatches as a table.
    swatches: list of (hex, name, description)
    """
    ncols = len(swatches)
    table = doc.add_table(rows=2, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)

    swatch_row = table.rows[0]
    label_row = table.rows[1]

    swatch_h_cm = 1.5
    col_w = 16.5 / ncols

    for i, (hex_color, name, desc) in enumerate(swatches):
        # Swatch cell
        sc = swatch_row.cells[i]
        set_cell_bg(sc, hex_color)
        set_column_width(sc, col_w)
        # Set height via XML
        tr = swatch_row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(swatch_h_cm * 567)))
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)
        # Add empty paragraph
        p = sc.paragraphs[0]
        p.add_run('')

        # Label cell
        lc = label_row.cells[i]
        set_cell_bg(lc, WHITE_HEX)
        set_column_width(lc, col_w)
        set_cell_margins(lc, top=40, bottom=40, left=60, right=60)
        p2 = lc.paragraphs[0]
        r1 = p2.add_run(name + '\n')
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9)
        r1.font.bold = True
        r1.font.color.rgb = DARK_TEXT
        r2 = p2.add_run('#' + hex_color)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(8)
        r2.font.color.rgb = MID_GREY

    return table


# ─── TOC Helper ──────────────────────────────────────────────────────────────

TOC_ENTRIES = []  # Will be populated as we add sections


def add_toc_entry(text, page_estimate, level=1):
    TOC_ENTRIES.append((text, page_estimate, level))


# ─── Document Assembly ────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page setup: A4
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK_TEXT

    # ── COVER PAGE ────────────────────────────────────────────────────────────

    # Top teal band — represented as a table spanning full width
    cover_band = doc.add_table(rows=1, cols=1)
    cover_band.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(cover_band)
    cb_cell = cover_band.cell(0, 0)
    set_cell_bg(cb_cell, TEAL_HEX)
    set_column_width(cb_cell, 16.5)
    set_cell_margins(cb_cell, top=400, bottom=400, left=300, right=300)
    tr = cover_band.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(4.0 * 567)))
    trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)
    # PHI wordmark
    p_phi = cb_cell.paragraphs[0]
    p_phi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_phi = p_phi.add_run('\u03c6')   # φ
    r_phi.font.name = 'Garamond'
    r_phi.font.size = Pt(72)
    r_phi.font.italic = True
    r_phi.font.color.rgb = WHITE

    # Spacer
    sp1 = doc.add_paragraph()
    set_spacing(sp1, before=0, after=0)

    # Brand name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p_name, before=200, after=60)
    r_phi2 = p_name.add_run('PHI')
    r_phi2.font.name = 'Garamond'
    r_phi2.font.size = Pt(42)
    r_phi2.font.italic = True
    r_phi2.font.bold = False
    r_phi2.font.color.rgb = TEAL
    r_fid = p_name.add_run('  FIDUCIAIRE')
    r_fid.font.name = 'Calibri'
    r_fid.font.size = Pt(20)
    r_fid.font.bold = False
    r_fid.font.color.rgb = TEAL

    # Tagline
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p_tag, before=60, after=60)
    r_tag = p_tag.add_run('\u201cLa rigueur suisse. L\u2019expertise genevoise.\u201d')
    r_tag.font.name = 'Garamond'
    r_tag.font.size = Pt(14)
    r_tag.font.italic = True
    r_tag.font.color.rgb = MID_GREY

    # Red divider line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p_div, before=160, after=160)
    set_para_border_bottom(p_div, RED_HEX, sz=8)

    # Document title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p_title, before=200, after=80)
    r_title = p_title.add_run('Brand Style Guide')
    r_title.font.name = 'Garamond'
    r_title.font.size = Pt(32)
    r_title.font.bold = False
    r_title.font.color.rgb = DARK_TEXT

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p_sub, before=40, after=40)
    r_sub = p_sub.add_run('Visual Identity & Design System v1.0')
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = MID_GREY

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p_date, before=40, after=40)
    r_date = p_date.add_run('Mars 2026')
    r_date.font.name = 'Calibri'
    r_date.font.size = Pt(11)
    r_date.font.color.rgb = MID_GREY

    # Spacer to push footer area
    for _ in range(6):
        sp = doc.add_paragraph()
        set_spacing(sp, before=0, after=0)

    # Footer info block
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p_foot, before=0, after=0)
    r_foot = p_foot.add_run('Gen\u00e8ve, Suisse  \u2022  Fiduciaire Premium  \u2022  Since 2010')
    r_foot.font.name = 'Calibri'
    r_foot.font.size = Pt(9)
    r_foot.font.color.rgb = MID_GREY

    # Bottom teal band
    bot_band = doc.add_table(rows=1, cols=1)
    bot_band.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(bot_band)
    bb_cell = bot_band.cell(0, 0)
    set_cell_bg(bb_cell, TEAL_DARK_HEX)
    set_column_width(bb_cell, 16.5)
    set_cell_margins(bb_cell, top=200, bottom=200, left=300, right=300)
    tr2 = bot_band.rows[0]._tr
    trPr2 = tr2.get_or_add_trPr()
    trH2 = OxmlElement('w:trHeight')
    trH2.set(qn('w:val'), str(int(1.2 * 567)))
    trH2.set(qn('w:hRule'), 'exact')
    trPr2.append(trH2)
    p_bb = bb_cell.paragraphs[0]
    p_bb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bb = p_bb.add_run('phi-fiduciaire.ch  \u2022  Gen\u00e8ve')
    r_bb.font.name = 'Calibri'
    r_bb.font.size = Pt(8.5)
    r_bb.font.color.rgb = WHITE

    # Page break after cover
    add_page_break(doc)

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────

    p_toc_title = doc.add_paragraph()
    set_spacing(p_toc_title, before=0, after=120)
    r_toc = p_toc_title.add_run('Table of Contents')
    r_toc.font.name = 'Garamond'
    r_toc.font.size = Pt(26)
    r_toc.font.bold = False
    r_toc.font.color.rgb = TEAL
    set_para_border_bottom(p_toc_title, TEAL_HEX, sz=4)

    toc_items = [
        ('1.  Brand Identity', 3),
        ('2.  Color Palette', 4),
        ('3.  Typography', 5),
        ('4.  Spacing System — Fibonacci / Golden Ratio', 7),
        ('5.  Design Principles', 8),
        ('6.  Image Treatment', 9),
        ('7.  UI Components', 11),
        ('8.  Section Visual Language', 13),
        ('9.  Applications — Carousels & GMB Posts', 14),
        ('10. Document Notes', 15),
    ]
    for title, page in toc_items:
        p_toc = doc.add_paragraph()
        set_spacing(p_toc, before=60, after=60)
        # Use tab stop for right-aligned page number
        pPr = p_toc._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab_el = OxmlElement('w:tab')
        tab_el.set(qn('w:val'), 'right')
        tab_el.set(qn('w:leader'), 'dot')
        tab_el.set(qn('w:pos'), '9072')  # ~16 cm in twips
        tabs.append(tab_el)
        pPr.append(tabs)
        r1 = p_toc.add_run(title)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = DARK_TEXT
        r2 = p_toc.add_run(f'\t{page}')
        r2.font.name = 'Calibri'
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = MID_GREY

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — BRAND IDENTITY
    # ══════════════════════════════════════════════════════════════════════════

    heading1(doc, '1.  Brand Identity', bm_id=1, bm_name='section1')

    make_table(doc,
        headers=['Attribute', 'Value'],
        rows=[
            ['Brand Name', 'PHI Fiduciaire'],
            ['Location', 'Gen\u00e8ve, Suisse'],
            ['Positioning', 'Premium Swiss fiduciary \u2014 rigorous, local, multilingual'],
            ['Primary Tagline', '\u201cLa rigueur suisse. L\u2019expertise genevoise.\u201d'],
            ['Brand Promise', '15 years of local expertise \u2022 Single point of contact \u2022 Response within 24h'],
            ['Languages Served', 'French, English, German, Polish'],
        ],
        col_widths_cm=[5.5, 11.0],
        stripe=True,
    )

    sp = doc.add_paragraph(); set_spacing(sp, 0, 0)

    heading2(doc, 'The \u03c6 Symbol')
    body(doc, 'PHI (\u03c6) is the golden ratio (1.618). The brand name is intentionally the mathematical symbol \u2014 it represents precision, proportion, and the Fibonacci system that underpins all spacing decisions.')
    note(doc, 'Critical: The \u03c6 symbol is NOT decorative; it is conceptual and must be respected in all communications. It connects the brand name directly to the mathematical principle of perfect proportion.')

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — COLOR PALETTE
    # ══════════════════════════════════════════════════════════════════════════

    heading1(doc, '2.  Color Palette', bm_id=2, bm_name='section2')

    body(doc, 'The PHI Fiduciaire palette is built around deep teal, drawing on Swiss precision and Geneva\'s lakeside environment. Every color has a defined, limited role.')

    sp = doc.add_paragraph(); set_spacing(sp, 0, 60)

    heading2(doc, 'Color Swatches')

    # Row 1: Primary + Dark Teal + Red
    add_color_swatch_row(doc, [
        ('1B3F3F', 'Primary Teal', 'Main brand color'),
        ('0F2626', 'Footer Dark', 'Hero/footer zones'),
        ('EF3A24', 'Accent Red', 'Sparingly — signal only'),
        ('FFFFFF', 'White', 'Backgrounds, reversed type'),
        ('F8F6F2', 'Off-white', 'Section backgrounds'),
    ])

    sp2 = doc.add_paragraph(); set_spacing(sp2, 0, 80)

    heading2(doc, 'Color Reference Table')

    make_table(doc,
        headers=['Name', 'Hex', 'RGB', 'Usage'],
        rows=[
            ['Primary Teal',   '#1b3f3f', 'rgb(27, 63, 63)',   'Main brand color, headings, nav, cards'],
            ['Footer Dark',    '#0f2626', 'rgb(15, 38, 38)',   'Footer, hero bottom band, contrast zones'],
            ['Accent Red',     '#ef3a24', 'rgb(239, 58, 36)',  'Sparingly: label underlines, dividers, hover states'],
            ['White',          '#ffffff', 'rgb(255, 255, 255)','Backgrounds, button text, reversed type'],
            ['Off-white',      '#f8f6f2', 'rgb(248, 246, 242)','Blog section, alternating backgrounds'],
            ['Border',         'rgba(27,63,63,0.12)', '\u2014', 'Subtle dividers, card borders'],
        ],
        col_widths_cm=[3.5, 2.5, 3.5, 7.0],
    )

    sp3 = doc.add_paragraph(); set_spacing(sp3, 0, 0)

    note(doc, 'Critical rule: The red accent (#ef3a24) must NEVER be overused. Maximum 2\u20133 instances visible at once in any single frame, carousel slide, or post. It is a signal, not a decoration.')

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — TYPOGRAPHY
    # ══════════════════════════════════════════════════════════════════════════

    heading1(doc, '3.  Typography', bm_id=3, bm_name='section3')

    body(doc, 'PHI Fiduciaire uses two typefaces: one editorial serif for display, one clean sans for body text. The contrast between these families IS the design \u2014 never substitute or mix additional typefaces.')

    heading2(doc, 'Display Font: Cormorant Garamond')

    make_table(doc,
        headers=['Attribute', 'Value'],
        rows=[
            ['Source', 'Google Fonts (open license)'],
            ['Used for', 'H1, H2, H3, pull quotes, italic emphasis lines'],
            ['Character', 'Elegant, editorial, Swiss-luxury feel'],
            ['Style note', 'Mix upright (H1) with italic (secondary headline) for typographic rhythm'],
        ],
        col_widths_cm=[4.5, 12.0],
    )

    sp = doc.add_paragraph(); set_spacing(sp, 0, 80)

    heading2(doc, 'Body Font: DM Sans')

    make_table(doc,
        headers=['Attribute', 'Value'],
        rows=[
            ['Source', 'Google Fonts (open license)'],
            ['Used for', 'Body text, labels, navigation, buttons, captions'],
            ['Character', 'Clean, modern, highly legible at small sizes'],
        ],
        col_widths_cm=[4.5, 12.0],
    )

    sp2 = doc.add_paragraph(); set_spacing(sp2, 0, 80)

    heading2(doc, 'Type Scale')

    make_table(doc,
        headers=['Element', 'Size', 'Weight', 'Font', 'Case'],
        rows=[
            ['H1 hero',         '72px / clamp(48px, 6vw, 72px)', '400',       'Cormorant Garamond', 'Sentence'],
            ['H1 italic line',  'Same as H1',                     '400 italic','Cormorant Garamond', 'Sentence'],
            ['H2 section',      'clamp(36px, 3.6vw, 52px)',        '400',       'Cormorant Garamond', 'Sentence'],
            ['H3 card title',   '20\u201322px',                    '500',       'DM Sans',             'Sentence'],
            ['Label / eyebrow', '11\u201312px',                    '500',       'DM Sans',             'UPPERCASE'],
            ['Body text',       '16\u201317px',                    '400',       'DM Sans',             'Sentence'],
            ['Button text',     '13px',                            '500',       'DM Sans',             'UPPERCASE, letter-spacing 0.08em'],
            ['Caption / meta',  '12\u201313px',                    '400',       'DM Sans',             'Sentence'],
            ['Nav links',       '13px',                            '400',       'DM Sans',             'Sentence'],
        ],
        col_widths_cm=[3.2, 4.5, 1.8, 4.0, 3.0],
    )

    sp3 = doc.add_paragraph(); set_spacing(sp3, 0, 80)

    heading2(doc, 'Label Style (Eyebrow Pattern)')
    body(doc, 'The eyebrow label is the consistent "section opener" pattern across the entire site. It must always appear in the same form:')
    bullet(doc, 'DM Sans, 11\u201312px, UPPERCASE, letter-spacing 0.12em')
    bullet(doc, 'Color: #1b3f3f on light backgrounds, #ffffff on dark backgrounds')
    bullet(doc, 'Followed immediately by: a 34px \u00d7 2px red (#ef3a24) horizontal line below the label')
    bullet(doc, 'Margin below: ~13px before the H2 headline')

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — SPACING SYSTEM
    # ══════════════════════════════════════════════════════════════════════════

    heading1(doc, '4.  Spacing System \u2014 Fibonacci / Golden Ratio', bm_id=4, bm_name='section4')

    body(doc, 'All spacing is derived from the Fibonacci sequence, starting at 13px. This creates a mathematically harmonious rhythm that mirrors the \u03c6 symbol itself.')

    heading2(doc, 'Spacing Scale')

    make_table(doc,
        headers=['Variable', 'Value', 'Use'],
        rows=[
            ['--space-xs',   '13px', 'Tight gaps, small padding'],
            ['--space-sm',   '21px', 'Component internal spacing'],
            ['--space-md',   '34px', 'Standard gap, card padding'],
            ['--space-lg',   '55px', 'Between components'],
            ['--space-xl',   '89px', 'Section padding'],
            ['--space-xxl',  '144px', 'Large section padding'],
            ['--space-hero', '233px', 'Hero-level vertical space'],
        ],
        col_widths_cm=[3.5, 2.5, 10.5],
    )

    sp = doc.add_paragraph(); set_spacing(sp, 0, 80)

    heading2(doc, 'Golden Ratio Nav / Band System')
    body(doc, 'All repeating bands use Fibonacci values: 34px and 55px alternating. This creates a palindrome rhythm that draws the eye to the center naturally.')

    make_table(doc,
        headers=['Element', 'Height', 'Background'],
        rows=[
            ['Mobile logo row',       '34px', '#0f2626 (Footer Dark)'],
            ['Mobile nav row',        '55px', 'White'],
            ['Hero bottom teal band', '55px', '#1b3f3f (Primary Teal)'],
            ['Hero bottom dark band', '34px', '#0f2626 (Footer Dark)'],
        ],
        col_widths_cm=[5.5, 2.5, 8.5],
    )

    sp2 = doc.add_paragraph(); set_spacing(sp2, 0, 0)
    note(doc, 'Palindrome rhythm: 34 / 55 / [photo] / 55 / 34 \u2014 the eye is drawn to the center naturally. Never break this proportional system.')

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — DESIGN PRINCIPLES
    # ══════════════════════════════════════════════════════════════════════════

    heading1(doc, '5.  Design Principles', bm_id=5, bm_name='section5')

    body(doc, 'These seven principles govern all visual decisions for PHI Fiduciaire. They are non-negotiable rules, not guidelines.')

    principles = [
        ('No Repeated Visual Patterns',
         'Each section must have a distinct visual treatment. Ghost numbers in Services \u2192 red accent lines in Why PHI \u2192 duotone images in Blog \u2192 etc. Repetition dilutes impact.'),
        ('Generous Whitespace',
         'Never crowd content. When in doubt, add more breathing room. Whitespace communicates luxury and confidence.'),
        ('Swiss Precision',
         'Every alignment is intentional. Left-aligned text blocks, controlled grid systems, nothing arbitrary. If something cannot be justified, it does not belong.'),
        ('Red is a Signal, Not Decoration',
         'Use the red accent maximum 2\u20133 times per viewport. It should feel like punctuation, not color. When red appears everywhere, it means nothing.'),
        ('Premium Feel Through Restraint',
         'Fewer elements, more space, higher quality. Avoid adding decorative elements that do not serve meaning. Subtract before you add.'),
        ('The \u03c6 Ghost Texture',
         'On dark sections, a giant semi-transparent "\u03c6" character (Cormorant Garamond, ~500px, 2\u20133% opacity) can appear as a background texture. Used in Final CTA only. Never overdone.'),
        ('Hierarchy Through Size Contrast',
         'Combine a large Cormorant Garamond display line with a much smaller DM Sans body \u2014 the contrast IS the design. Scale difference should be dramatic, not subtle.'),
    ]

    for i, (title, desc) in enumerate(principles, 1):
        p_num = doc.add_paragraph()
        set_spacing(p_num, before=120, after=40)
        r_n = p_num.add_run(f'{i}.')
        r_n.font.name = 'Garamond'
        r_n.font.size = Pt(18)
        r_n.font.bold = True
        r_n.font.color.rgb = TEAL
        r_space = p_num.add_run('  ')
        r_space.font.size = Pt(11)
        r_title2 = p_num.add_run(title)
        r_title2.font.name = 'Calibri'
        r_title2.font.size = Pt(11)
        r_title2.font.bold = True
        r_title2.font.color.rgb = DARK_TEXT
        p_desc = doc.add_paragraph()
        set_spacing(p_desc, before=20, after=60)
        pPr = p_desc._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '360')
        pPr.append(ind)
        r_desc = p_desc.add_run(desc)
        r_desc.font.name = 'Calibri'
        r_desc.font.size = Pt(10.5)
        r_desc.font.color.rgb = DARK_TEXT

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — IMAGE TREATMENT
    # ══════════════════════════════════════════════════════════════════════════

    heading1(doc, '6.  Image Treatment', bm_id=6, bm_name='section6')

    heading2(doc, 'Photography Style')
    bullet(doc, 'Subject: Architectural / geometric \u2014 spiral staircases, Geneva lakeside, Swiss precision imagery')
    bullet(doc, 'Color palette of photos: neutral, cool, with warm highlights acceptable')
    bullet(doc, 'NO stock photo clich\u00e9s (no handshakes, no generic office shots)')
    bullet(doc, 'PHI uses its own photography (phi-stairs series) \u2014 always prefer proprietary images')

    sp = doc.add_paragraph(); set_spacing(sp, 0, 60)

    heading2(doc, 'Overlay Technique \u2014 Desktop (Left-to-Right)')
    body(doc, 'Used on the hero section. The goal is to keep the photo vivid on the right while creating a dark, legible text zone on the left.')

    make_table(doc,
        headers=['Layer', 'Rule'],
        rows=[
            ['Source photo', 'Original color photo \u2014 NO desaturation, NO duotone'],
            ['Overlay 1', 'Left solid teal (#1b3f3f) covering 15\u201320% of width \u2192 hard cut \u2192 smooth gradient fading to transparent at ~75%'],
            ['Overlay 2', 'Bottom dark teal fade: bottom 25% fades from transparent to near-solid #0f2626'],
            ['Result', 'Photo vivid and clear on right; text zone dark and legible on left'],
        ],
        col_widths_cm=[3.0, 13.5],
    )

    sp2 = doc.add_paragraph(); set_spacing(sp2, 0, 80)

    heading2(doc, 'Overlay Technique \u2014 Mobile (Top-to-Bottom)')

    make_table(doc,
        headers=['Zone', 'Treatment'],
        rows=[
            ['0\u201350%',  'Pure photo, no overlay'],
            ['50\u201380%', 'Gradient fading in to 95% teal (#1b3f3f) opacity'],
            ['80%',         'Hard cut \u2192 solid main teal (#1b3f3f)'],
            ['90%',         'Hard cut \u2192 solid footer dark (#0f2626)'],
        ],
        col_widths_cm=[3.0, 13.5],
    )

    sp3 = doc.add_paragraph(); set_spacing(sp3, 0, 80)

    heading2(doc, 'Duotone Technique (Blog Cards, Quote Sections)')

    make_table(doc,
        headers=['Step', 'Action'],
        rows=[
            ['Step 1', 'Apply filter: grayscale(60%) brightness(1.1) to image'],
            ['Step 2', 'Place a teal gradient overlay with mix-blend-mode: color at 50% opacity'],
            ['Gradient', 'linear-gradient(150deg, #2a6565 0%, #1b3f3f 100%)'],
            ['Result', 'Teal-tinted image that integrates seamlessly with brand palette'],
        ],
        col_widths_cm=[2.5, 14.0],
    )

    sp4 = doc.add_paragraph(); set_spacing(sp4, 0, 80)

    heading2(doc, 'Hard Cut Technique')
    body(doc, 'A "hard cut" is a pixel-sharp color band transition achieved by placing two CSS color stops at the same percentage:')

    p_code = doc.add_paragraph()
    set_spacing(p_code, before=60, after=60)
    pPr = p_code._p.get_or_add_pPr()
    ind_el = OxmlElement('w:ind')
    ind_el.set(qn('w:left'), '400')
    pPr.append(ind_el)
    p_bg = OxmlElement('w:pBdr')
    l_b = OxmlElement('w:left')
    l_b.set(qn('w:val'), 'single')
    l_b.set(qn('w:sz'), '8')
    l_b.set(qn('w:space'), '6')
    l_b.set(qn('w:color'), TEAL_HEX)
    p_bg.append(l_b)
    pPr.append(p_bg)
    r_code = p_code.add_run('rgba(27,63,63,1.00) 80%, rgba(15,38,38,1.00) 80%')
    r_code.font.name = 'Courier New'
    r_code.font.size = Pt(9.5)
    r_code.font.color.rgb = TEAL

    body(doc, 'Used to create distinct color zones echoing the nav band structure.')

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — UI COMPONENTS
    # ══════════════════════════════════════════════════════════════════════════

    heading1(doc, '7.  UI Components', bm_id=7, bm_name='section7')

    # Section Label
    heading2(doc, 'Section Label (Eyebrow)')
    make_table(doc,
        headers=['Property', 'Value'],
        rows=[
            ['Font', 'DM Sans, 11px, uppercase, letter-spacing 0.12em'],
            ['Color', '#1b3f3f (white on dark backgrounds)'],
            ['Accent element', '34px \u00d7 2px red (#ef3a24) horizontal line below the label'],
            ['Margin below', '~13px before H2'],
        ],
        col_widths_cm=[4.0, 12.5],
    )
    sp = doc.add_paragraph(); set_spacing(sp, 0, 80)

    # Buttons
    heading2(doc, 'Buttons')
    make_table(doc,
        headers=['Variant', 'Background', 'Text', 'Border', 'Radius', 'Padding', 'Hover'],
        rows=[
            ['Primary (dark bg)',    'White',        'Teal',           'None',              '50px (pill)', '16px \u00d7 34px', 'Red bg + red glow shadow'],
            ['Ghost / Secondary',    'Transparent',  'White at 55%',   '1.5px white',       '50px (pill)', '16px \u00d7 34px', 'Border brightens to 100%'],
            ['Outline (light bg)',   'Transparent',  'Teal',           '1px teal',          'None (rect)', '14px \u00d7 28px', 'Teal fill'],
        ],
        col_widths_cm=[3.2, 2.4, 2.4, 2.4, 2.2, 2.2, 1.7],
    )
    sp2 = doc.add_paragraph(); set_spacing(sp2, 0, 0)
    note(doc, 'All button text: 13px DM Sans, UPPERCASE, letter-spacing 0.08em. Never use sentence case on buttons.')
    sp3 = doc.add_paragraph(); set_spacing(sp3, 0, 80)

    # Cards
    heading2(doc, 'Cards (Why PHI, Blog)')
    make_table(doc,
        headers=['Property', 'Value'],
        rows=[
            ['Border top',    '1px solid rgba(27,63,63,0.12) + 34px \u00d7 2px red accent line above'],
            ['Background',    'White (or off-white #f8f6f2 on off-white sections = white card)'],
            ['Drop shadow',   'None \u2014 border only'],
            ['Link style',    'Uppercase 11px DM Sans with \u2192 arrow, teal color, border-bottom on hover'],
        ],
        col_widths_cm=[4.0, 12.5],
    )
    sp4 = doc.add_paragraph(); set_spacing(sp4, 0, 80)

    # FAQ
    heading2(doc, 'FAQ Accordion')
    make_table(doc,
        headers=['Property', 'Value'],
        rows=[
            ['Toggle icon',      'Bare +/\u2212 character (no circle, no box) \u2014 22px, weight 300'],
            ['Closed state',     'Normal weight question text, no border'],
            ['Open state',       '2px red (#ef3a24) left border on entire item'],
            ['Open state text',  'Question text weight increases to 500'],
            ['Animation',        'Max-height transition for smooth open/close'],
        ],
        col_widths_cm=[4.0, 12.5],
    )
    sp5 = doc.add_paragraph(); set_spacing(sp5, 0, 80)

    # Navigation Desktop
    heading2(doc, 'Navigation \u2014 Desktop')
    make_table(doc,
        headers=['Property', 'Value'],
        rows=[
            ['Behavior',     'Fixed, transparent on hero \u2192 solid teal (#1b3f3f) on scroll'],
            ['Logo',         '"PHI" in Cormorant Garamond italic + "FIDUCIAIRE" in DM Sans small caps'],
            ['Nav links',    '13px DM Sans, white at 75% opacity, sentence case'],
            ['CTA button',   'Outlined, white border, ghost style'],
            ['Phone',        'Visible on desktop nav'],
        ],
        col_widths_cm=[3.5, 13.0],
    )
    sp6 = doc.add_paragraph(); set_spacing(sp6, 0, 80)

    # Navigation Mobile
    heading2(doc, 'Navigation \u2014 Mobile')
    make_table(doc,
        headers=['Band', 'Height', 'Background', 'Content'],
        rows=[
            ['Band 1 (top)',    '34px', '#0f2626 (Footer Dark)', 'Centered logo wordmark'],
            ['Band 2 (bottom)', '55px', 'White',                 'Hamburger (left) + CTA button (right)'],
        ],
        col_widths_cm=[3.2, 2.0, 4.8, 6.5],
    )
    sp7 = doc.add_paragraph(); set_spacing(sp7, 0, 0)
    note(doc, 'CTA button icon: User/person icon in red (#ef3a24), thin stroke 1.5px. The two-band mobile nav mirrors the Fibonacci band rhythm (34 / 55).')

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 8 — SECTION VISUAL LANGUAGE
    # ══════════════════════════════════════════════════════════════════════════

    heading1(doc, '8.  Section Visual Language', bm_id=8, bm_name='section8')

    body(doc, 'Each page section has a distinct visual identity. No two sections should look the same. Below is the master reference for section treatments.')

    make_table(doc,
        headers=['Section', 'Background', 'Key Visual', 'Notes'],
        rows=[
            ['Hero',          'Full-bleed photo + directional overlay', 'Spiral staircase or architectural',    'Left/bottom gradient; text bottom-left'],
            ['Services',      'White',                                   'Ghost \u03c6 numbers behind cards',       '3-column grid'],
            ['Why PHI',       'White',                                   '40/60 header split + 3-col cards',    'Red accent line per card'],
            ['Pricing Teaser','Full-bleed water photo',                  'Left solid dark band (15%) + gradient','Mirabaud-style treatment'],
            ['Local SEO',     'White',                                   'Left: rich text; Right: teal map + NAP','2-column layout'],
            ['FAQ',           'White',                                   'Accordion; red left border on open',  'No decorative elements'],
            ['Blog',          'Off-white (#f8f6f2)',                     'White cards; duotone images',         '3-col grid, 34px gap'],
            ['Final CTA',     'Full-bleed landscape photo',              'Ghost \u03c6; left-gradient dark overlay',  'Single white CTA button'],
            ['Footer',        'Footer dark (#0f2626)',                   '4-column layout',                     'White text on dark'],
        ],
        col_widths_cm=[3.0, 3.5, 5.0, 5.0],
    )

    sp = doc.add_paragraph(); set_spacing(sp, 0, 80)
    note(doc, 'The ghost \u03c6 texture (Cormorant Garamond, ~500px, 2\u20133% opacity on dark sections) is reserved exclusively for the Final CTA section. Never apply it to more than one section per page.')

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 9 — APPLICATIONS
    # ══════════════════════════════════════════════════════════════════════════

    heading1(doc, '9.  Applications \u2014 Carousels & GMB Posts', bm_id=9, bm_name='section9')

    heading2(doc, 'Carousel Slides')

    make_table(doc,
        headers=['Property', 'Specification'],
        rows=[
            ['Format',       '1080\u00d71080px (square) or 1080\u00d71350px (portrait)'],
            ['Safe zone',    '80px margin on all sides'],
            ['Typography',   'Cormorant Garamond for headline; DM Sans for body/caption'],
            ['Red accent',   'Max 1\u20132 uses per slide (label line OR small divider, not both)'],
            ['Slide structure','Label (eyebrow) \u2192 Headline \u2192 Body \u2192 CTA or stat'],
            ['Consistency',  'First and last slide must use brand teal (#1b3f3f) background'],
        ],
        col_widths_cm=[4.0, 12.5],
    )

    heading3(doc, 'Background Options')
    bullet(doc, 'Solid #1b3f3f with ghost \u03c6 texture')
    bullet(doc, 'Photo with left-to-right or bottom overlay (same technique as hero)')
    bullet(doc, 'Off-white #f8f6f2 for light slides')

    sp = doc.add_paragraph(); set_spacing(sp, 0, 80)

    heading2(doc, 'GMB (Google Business Profile) Posts')

    make_table(doc,
        headers=['Property', 'Specification'],
        rows=[
            ['Format',         '1200\u00d7900px landscape or 1080\u00d71080px square'],
            ['Wordmark',       'PHI Fiduciaire wordmark: bottom-right or bottom-center'],
            ['Tone',           'Professional, informative, local \u2014 never salesy'],
            ['Image',          'Brand photography or teal-overlaid architectural photos'],
            ['Text on image',  'Maximum 20% of image area (Google guideline)'],
            ['Color lead',     'Teal (#1b3f3f) primary; red (#ef3a24) sparingly as accent'],
        ],
        col_widths_cm=[4.0, 12.5],
    )

    sp2 = doc.add_paragraph(); set_spacing(sp2, 0, 80)

    heading2(doc, 'Social Media General Rules')

    make_table(doc,
        headers=['Rule', 'Detail'],
        rows=[
            ['Typography',  'Never use fonts other than Cormorant Garamond + DM Sans'],
            ['Color',       'Never use colors outside the brand palette'],
            ['\u03c6 presence', 'Always include the \u03c6 in some form: wordmark, ghost texture, or explicit reference'],
            ['Whitespace',  'Maintain generous whitespace principle \u2014 never crowd the frame'],
        ],
        col_widths_cm=[3.5, 13.0],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 10 — DOCUMENT NOTES
    # ══════════════════════════════════════════════════════════════════════════

    heading1(doc, '10. Document Notes', bm_id=10, bm_name='section10')

    heading2(doc, 'Version Information')
    make_table(doc,
        headers=['Field', 'Value'],
        rows=[
            ['Document version', 'v1.0'],
            ['Date',             'Mars 2026'],
            ['Status',           'Living brief \u2014 will be updated as visual identity evolves'],
            ['Primary reference implementation', 'https://noxuulab.github.io/phi-fiduciaire-design/'],
        ],
        col_widths_cm=[5.0, 11.5],
    )

    sp = doc.add_paragraph(); set_spacing(sp, 0, 80)

    heading2(doc, 'Usage Policy')
    bullet(doc, 'This document is the single source of truth for all PHI Fiduciaire visual communications.')
    bullet(doc, 'Any deviation from these guidelines requires written approval.')
    bullet(doc, 'Third-party vendors, designers, and agencies must receive this guide before beginning any work.')
    bullet(doc, 'Digital assets (logo files, color palettes, font licenses) are maintained separately and available upon request.')

    sp2 = doc.add_paragraph(); set_spacing(sp2, 0, 80)

    # Closing statement with teal background
    close_band = doc.add_table(rows=1, cols=1)
    close_band.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(close_band)
    cc = close_band.cell(0, 0)
    set_cell_bg(cc, TEAL_HEX)
    set_column_width(cc, 16.5)
    set_cell_margins(cc, top=400, bottom=400, left=500, right=500)

    p_c1 = cc.paragraphs[0]
    p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c1 = p_c1.add_run('PHI Fiduciaire')
    r_c1.font.name = 'Garamond'
    r_c1.font.size = Pt(20)
    r_c1.font.italic = True
    r_c1.font.color.rgb = WHITE

    p_c2 = cc.add_paragraph()
    p_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c2 = p_c2.add_run('\u201cLa rigueur suisse. L\u2019expertise genevoise.\u201d')
    r_c2.font.name = 'Garamond'
    r_c2.font.size = Pt(12)
    r_c2.font.italic = True
    r_c2.font.color.rgb = RGBColor(0xcc, 0xdd, 0xdd)

    p_c3 = cc.add_paragraph()
    p_c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c3 = p_c3.add_run('Brand Style Guide v1.0 \u2014 Mars 2026 \u2014 Strictly Confidential')
    r_c3.font.name = 'Calibri'
    r_c3.font.size = Pt(8.5)
    r_c3.font.color.rgb = RGBColor(0x99, 0xbb, 0xbb)

    return doc


def main():
    print("Building PHI Fiduciaire Brand Style Guide...")
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")
    size = os.path.getsize(OUTPUT_PATH)
    print(f"File size: {size:,} bytes ({size/1024:.1f} KB)")


if __name__ == '__main__':
    main()
